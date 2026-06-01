import io
import os
import re
import base64
import zlib
import requests
import json
import mistune
from docx import Document
from docx.shared import Inches
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from typing import List, Dict, Any
import time

def render_mermaid_to_image(mermaid_code: str) -> io.BytesIO:
    """Fetch PNG from kroki.io using POST method for large payloads"""
    url = "https://kroki.io/mermaid/png"
    headers = {
        "Content-Type": "text/plain",
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
    }
    for attempt in range(3):
        try:
            response = requests.post(url, data=mermaid_code.strip().encode('utf-8'), headers=headers, timeout=15)
            response.raise_for_status()
            return io.BytesIO(response.content)
        except Exception as e:
            error_msg = str(e)
            if hasattr(e, 'response') and e.response is not None:
                error_msg = e.response.text
            print(f"Failed to render mermaid (attempt {attempt + 1}): {error_msg}")
            if attempt < 2:
                time.sleep(1)
    return None

def clear_document(doc: Document):
    for paragraph in doc.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)
        paragraph._p = paragraph._element = None
    for table in doc.tables:
        tbl = table._element
        tbl.getparent().remove(tbl)
        table._tbl = table._element = None

def set_table_width_100(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000') # 5000 pct = 100% width
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)

def render_inline(runs_container, children, bold=False, italic=False):
    if not children:
        return
    for child in children:
        ctype = child.get("type")
        if ctype == "text":
            run = runs_container.add_run(child.get("raw", ""))
            if bold: run.bold = True
            if italic: run.italic = True
        elif ctype == "strong":
            render_inline(runs_container, child.get("children", []), bold=True, italic=italic)
        elif ctype == "emphasis":
            render_inline(runs_container, child.get("children", []), bold=bold, italic=True)
        elif ctype == "codespan":
            run = runs_container.add_run(child.get("raw", ""))
            if bold: run.bold = True
            if italic: run.italic = True
        else:
            if "raw" in child:
                run = runs_container.add_run(child.get("raw", ""))
            elif "children" in child:
                render_inline(runs_container, child.get("children", []), bold=bold, italic=italic)

def render_ast_to_docx(doc, ast_nodes):
    for node in ast_nodes:
        ctype = node.get("type")
        
        if ctype == "paragraph":
            p = doc.add_paragraph()
            render_inline(p, node.get("children", []))
            
        elif ctype == "heading":
            # 降级渲染正文中的 Markdown 标题，防止污染 Word 官方 TOC 目录
            p = doc.add_paragraph()
            render_inline(p, node.get("children", []), bold=True)
            
        elif ctype == "list":
            children = node.get("children", [])
            is_ordered = node.get("attrs", {}).get("ordered", False)
            for li in children:
                if li.get("type") == "list_item":
                    style = 'List Number' if is_ordered else 'List Bullet'
                    try:
                        p = doc.add_paragraph(style=style)
                    except KeyError:
                        p = doc.add_paragraph()
                    li_children = li.get("children", [])
                    for lc in li_children:
                        if lc.get("type") == "block_text":
                            render_inline(p, lc.get("children", []))
                        
        elif ctype == "table":
            table_children = node.get("children", [])
            all_rows = []
            for child in table_children:
                child_type = child.get("type")
                if child_type == "table_head":
                    cells = []
                    for tc in child.get("children", []):
                        if tc.get("type") == "table_cell":
                            cells.append(tc.get("children", []))
                    if cells:
                        all_rows.append(cells)
                elif child_type == "table_body":
                    for tr in child.get("children", []):
                        if tr.get("type") == "table_row":
                            cells = []
                            for tc in tr.get("children", []):
                                if tc.get("type") == "table_cell":
                                    cells.append(tc.get("children", []))
                            if cells:
                                all_rows.append(cells)
            
            if all_rows and all_rows[0]:
                try:
                    table = doc.add_table(rows=len(all_rows), cols=len(all_rows[0]))
                    try:
                        table.style = 'Table Grid'
                    except KeyError:
                        pass
                    set_table_width_100(table)
                    for r_idx, row_data in enumerate(all_rows):
                        for c_idx, cell_ast in enumerate(row_data):
                            if c_idx < len(table.columns):
                                cell = table.cell(r_idx, c_idx)
                                cell.text = "" # Clear default paragraph text
                                p = cell.paragraphs[0]
                                render_inline(p, cell_ast)
                except Exception as e:
                    print(f"Table generation error: {e}")
                    
        elif ctype == "block_code":
            info = node.get("attrs", {}).get("info", "")
            val = node.get("raw", "")
            if info and info.strip().lower() == "mermaid":
                img_stream = render_mermaid_to_image(val)
                if img_stream:
                    doc.add_picture(img_stream, width=Inches(5.0))
                else:
                    doc.add_paragraph(f"[图表渲染失败: 大模型生成的 Mermaid 语法错误或网络异常]\n```mermaid\n{val}\n```")
            else:
                doc.add_paragraph(val)

def rebuild_docx_from_outline(outline: List[Dict], template_path: str = None) -> io.BytesIO:
    if template_path and os.path.exists(template_path):
        try:
            doc = Document(template_path)
            clear_document(doc)
        except Exception as e:
            print(f"Failed to load template {template_path}: {e}")
            doc = Document()
    else:
        doc = Document()
    
    md_parser = mistune.create_markdown(renderer='ast', plugins=['table'])
    
    for item in outline:
        title = item.get("title", "")
        level = item.get("level", 1)
        content = item.get("content", "")
        
        if level < 1: level = 1
        elif level > 9: level = 9
            
        style_name = f"Heading {level}"
        
        try:
            doc.add_paragraph(title, style=style_name)
        except KeyError:
            doc.add_paragraph(title)
            
        if content:

            ast_nodes = md_parser(content)
            render_ast_to_docx(doc, ast_nodes)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer
