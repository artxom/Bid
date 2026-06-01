import io
from docx import Document
from services.docx_builder import render_mermaid_to_image

doc = Document()
code1 = "graph TD; A-->B;"
code2 = "graph TD; C-->D;"

img1 = render_mermaid_to_image(code1)
if img1:
    print("img1 success")
    doc.add_picture(img1)
else:
    print("img1 failed")

img2 = render_mermaid_to_image(code2)
if img2:
    print("img2 success")
    doc.add_picture(img2)
else:
    print("img2 failed")

doc.save("test_multi.docx")
print("Done")
