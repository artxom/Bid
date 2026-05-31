import pytest
from fastapi.testclient import TestClient
from main import app
from docx import Document
import io

client = TestClient(app)

def create_mock_docx():
    doc = Document()
    doc.add_heading('第一章 项目背景', level=1)
    doc.add_paragraph('这是正文内容。')
    doc.add_heading('1.1 建设目标', level=2)
    doc.add_paragraph('更多正文。')
    doc.add_heading('第二章 技术方案', level=1)
    
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def test_upload_and_parse_outline():
    file_stream = create_mock_docx()
    files = {"file": ("test.docx", file_stream, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    response = client.post("/upload", files=files)
    
    assert response.status_code == 200
    data = response.json()
    assert "outline" in data
    outline = data["outline"]
    
    # 验证提取的大纲
    assert len(outline) == 3
    assert outline[0]["title"] == "第一章 项目背景"
    assert outline[0]["level"] == 1
    assert outline[1]["title"] == "1.1 建设目标"
    assert outline[1]["level"] == 2
    assert outline[2]["title"] == "第二章 技术方案"
    assert outline[2]["level"] == 1
