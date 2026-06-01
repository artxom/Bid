import io
import pytest
from unittest.mock import patch, MagicMock
from docx import Document
from services.docx_builder import render_mermaid_to_image, rebuild_docx_from_outline

def test_render_mermaid_to_image_success():
    with patch('services.docx_builder.requests.post') as mock_post:
        mock_response = MagicMock()
        mock_response.content = b"fake_png_data"
        mock_post.return_value = mock_response
        
        mermaid_code = "graph TD; A-->B;"
        result = render_mermaid_to_image(mermaid_code)
        
        assert result is not None
        assert result.read() == b"fake_png_data"
        mock_post.assert_called_once()

def test_render_mermaid_to_image_failure():
    with patch('services.docx_builder.requests.post') as mock_post:
        mock_post.side_effect = Exception("Connection Error")
        
        with patch('services.docx_builder.time.sleep') as mock_sleep:
            result = render_mermaid_to_image("graph TD; A-->B;")
            
            assert result is None
            assert mock_post.call_count == 3
            assert mock_sleep.call_count == 2

def test_rebuild_docx_from_outline_with_table():
    outline = [
        {
            "id": "node-1",
            "title": "表格测试章节",
            "level": 1,
            "content": "下面是一个测试表格\n\n| 表头1 | 表头2 |\n|---|---|\n| 数据1 | 数据2 |\n| 数据3 | 数据4 |\n"
        }
    ]
    
    buffer = rebuild_docx_from_outline(outline)
    assert buffer is not None
    
    buffer.seek(0)
    doc = Document(buffer)
    
    assert len(doc.tables) == 1
    table = doc.tables[0]
    
    assert len(table.rows) == 3
    assert len(table.columns) == 2
    
    assert "表头1" in table.cell(0, 0).text
    assert "表头2" in table.cell(0, 1).text
    assert "数据1" in table.cell(1, 0).text
    assert "数据4" in table.cell(2, 1).text
